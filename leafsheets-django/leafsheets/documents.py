"""

--- LeafSheets ---

Serialzers

Created on February 13, 2019
~ satyameva_jayate
"""

# Imports 

import logging
import random
import string 
import json
import re

import docx

# Log

logger = logging.getLogger(__name__)

# Helpers

def calc_var_count(var_dict):
    var_count = 0
    context = var_dict.get('context', None)
    if context:
        for item in context:
            if item.get('type', None) == 'VARIABLE':
                var_count += 1
    return var_count

def calc_provided_input_counts(var_dict):
    provided_input_count = 0
    provided_required_input_count = 0
    context = var_dict.get('context', None)
    if context:
        for item in context:
            if item.get('type', None) == 'INPUT':
                replace = item.get('replace', None)
                if replace != None and len(replace) > 0:
                    provided_input_count += 1
                    if item.get('meta').get('required') is True:
                        provided_required_input_count += 1
    return provided_input_count, provided_required_input_count

def calc_input_counts(var_dict):
    total_input_count = 0
    required_input_count = 0
    context = var_dict.get('context', None)
    if context:
        for item in context:
            if item.get('type', None) == 'INPUT':
                total_input_count += 1
                if item.get('meta').get('required') is True:
                    required_input_count += 1
    return total_input_count, required_input_count

def var_strings_from_doc(doc):
    """Returns the variable dictionary from the template document.
    """
    matches = []
    # Parse the tables for matching strings.
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    matches.extend(re.findall(r'\{\%(.+?)\%\}', paragraph.text))
    # Parse the paragraphs for matching strings.
    for paragraph in  doc.paragraphs:
        matches.extend(re.findall(r'\{\%(.+?)\%\}', paragraph.text))
    return matches

def match_dict_meta_from_match(kind, splits):
    if kind == 'VARIABLE':
        return {
            'name': splits[1],
        }
    else: 
        # Previous Style
        try:
            value = True if splits[4] == 'TRUE' else False
            inherit_previous_style = value
        except IndexError:
            inherit_previous_style = False
        # Group
        try:
            group = splits[5]
        except IndexError:
            group = None
        # Required
        try:
            required = True if splits[6] == 'REQUIRED' else False
        except IndexError:
            required = False
        # Default
        try:
            default = splits[7]
        except IndexError:
            default = None
        return {
            'name': splits[2],
            'prompt': splits[3],
            'inherit_previous_style': inherit_previous_style,
            'group': group,
            'required': required,
            'default': default
        }

def match_dict_from_match(match):
    splits = match.split('|')
    splits = [split.strip() for split in splits]
    return {
        'match': match,
        'type': splits[0].strip(),
        'meta': match_dict_meta_from_match(splits[0].strip(), splits),
        'find': ''.join(random.choices(string.ascii_uppercase + string.digits, k=6)),
        'replace': None,
    }

def variable_dict_from_doc(doc):
    """Return a JSON dict constructed from the regex matchs 
    within a template document.
    """
    result = {'context': []}
    context = result['context']
    matches = var_strings_from_doc(doc)
    for match in matches:
        match_dict = match_dict_from_match(match)
        context.append(match_dict)
    return result

def new_doc_from_template_with_var_substitution(var_dict, doc):
    """Returns a new document with the variables subsituted.
    """
    for var in var_dict['context']:
        old = '{{% {} %}}'.format(var['match'].strip())
        new = '{{% {} %}}'.format(var['find'])
        docx_find_replace_text(doc, old, new)
    return doc

def save_doc_with_user_var_substitution(var_dict, doc):
    """Returns a new document with the variables subsituted.
    """
    for var in var_dict['context']:
        old = '{{% {} %}}'.format(var['find'].strip())
        new = var['replace']
        default = var['meta'].get('default', None)
        if new is not None:
            docx_find_replace_text(doc, old, new)
        else:
            if default is not None:
                docx_find_replace_text(doc, old, default)
            else:
                docx_find_replace_text(doc, old, '--INPUT NOT PROVIDED--'.format(var['meta']['name']))
    return doc

def docx_find_replace_text(doc, search_text, replace_text):
    paragraphs = list(doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraphs.append(paragraph)
    for p in paragraphs:
        if search_text in p.text:
            inline = p.runs
            # Replace strings and retain the same style.
            # The text to be replaced can be split over several runs so
            # search through, identify which runs need to have text replaced
            # then replace the text in those identified
            started = False
            search_index = 0
            # found_runs is a list of (inline index, index of match, length of match)
            found_runs = list()
            found_all = False
            replace_done = False
            for i in range(len(inline)):
                # case 1: found in single run so short circuit the replace
                if search_text in inline[i].text and not started:
                    found_runs.append((i, inline[i].text.find(search_text), len(search_text)))
                    text = inline[i].text.replace(search_text, str(replace_text))
                    inline[i].text = text
                    replace_done = True
                    found_all = True
                    break
                if search_text[search_index] not in inline[i].text and not started:
                    # keep looking ...
                    continue
                # case 2: search for partial text, find first run
                if search_text[search_index] in inline[i].text and inline[i].text[-1] in search_text and not started:
                    # check sequence
                    start_index = inline[i].text.find(search_text[search_index])
                    check_length = len(inline[i].text)
                    for text_index in range(start_index, check_length):
                        if inline[i].text[text_index] != search_text[search_index]:
                            # no match so must be false positive
                            break
                    if search_index == 0:
                        started = True
                    chars_found = check_length - start_index
                    search_index += chars_found
                    found_runs.append((i, start_index, chars_found))
                    if search_index != len(search_text):
                        continue
                    else:
                        # found all chars in search_text
                        found_all = True
                        break
                # case 2: search for partial text, find subsequent run
                if search_text[search_index] in inline[i].text and started and not found_all:
                    # check sequence
                    chars_found = 0
                    check_length = len(inline[i].text)
                    for text_index in range(0, check_length):
                        if inline[i].text[text_index] == search_text[search_index]:
                            search_index += 1
                            chars_found += 1
                        else:
                            break
                    # no match so must be end
                    found_runs.append((i, 0, chars_found))
                    if search_index == len(search_text):
                        found_all = True
                        break
            if found_all and not replace_done:
                for i, item in enumerate(found_runs):
                    index, start, length = [t for t in item]
                    if i == 0:
                        text = inline[index].text.replace(inline[index].text[start:start + length], str(replace_text))
                        inline[index].text = text
                    else:
                        text = inline[index].text.replace(inline[index].text[start:start + length], '')
                        inline[index].text = text