"""

--- LeafSheets ---

Models

Created on Wednesday, Oct 2, 2019
~ satyameva_jayate
"""

# Imports

import re
import os
import json
import uuid
import requests
from copy import deepcopy

from django.db import models
from django.utils import timezone
from django.core.files import File
from django.utils.safestring import mark_safe
from django.contrib.auth import get_user_model
from django.db.models.signals import post_save, pre_save
from django.contrib.postgres.fields import ArrayField, JSONField

import docx

import shortuuid

from imagekit.models import ProcessedImageField
from imagekit.processors import ResizeToFill

from leafsheets_django.utils import PathAndRename, unique_slug_generator
from leafsheets_django.storage_backends import PrivateMediaStorage

from leafsheets.documents import (
    variable_dict_from_doc,
    new_doc_from_template_with_var_substitution,
    save_doc_with_user_var_substitution
    )
from leafsheets.managers import (
    UserSheetManager,
    DocumentManager,
    CompanyManager,
    SheetManager
    )

from commerce.models import Item, Charge

# User

User = get_user_model()

# CHOICES

NEW_CHOICES = (
    ('N', 'New'),
    ('E', 'Existing'),
)

CATEGORY_CHOICES = (
    ('M', 'Medical'),
    ('R', 'Recreational'),
    ('B', 'Both'),
)

INC_CHOICES = (
    ('C', 'C Corp'),
    ('S', 'S Corp'),
    ('LLC', 'Limited Liability Corp'),
    ('DBA', 'Doing Business As')
)

USE_CASE_CHOICES = (
    ('P', 'Production'),
    ('C', 'Cultivation'),
    ('D', 'Dispensary'),
    ('M', 'Multiple Purposes')
)

STATE_CHOICES = (
    ('AL','Alabama'),
    ('AK','Alaska'),
    ('AZ','Arizona'),
    ('AR','Arkansas'),
    ('CA','California'),
    ('CO','Colorado'),
    ('CT','Connecticut'),
    ('DE','Delaware'),
    ('FL','Florida'),
    ('GA','Georgia'),
    ('HI','Hawaii'),
    ('ID','Idaho'),
    ('IL','Illinois'),
    ('IN','Indiana'),
    ('IA','Iowa'),
    ('KS','Kansas'),
    ('KY','Kentucky'),
    ('LA','Louisiana'),
    ('ME','Maine'),
    ('MD','Maryland'),
    ('MA','Massachusetts'),
    ('MI','Michigan'),
    ('MN','Minnesota'),
    ('MS','Mississippi'),
    ('MO','Missouri'),
    ('MT','Montana'),
    ('NE','Nebraska'),
    ('NV','Nevada'),
    ('NH','New Hampshire'),
    ('NJ','New Jersey'),
    ('NM','New Mexico'),
    ('NY','New York'),
    ('NC','North Carolina'),
    ('ND','North Dakota'),
    ('OH','Ohio'),
    ('OK','Oklahoma'),
    ('OR','Oregon'),
    ('PA','Pennsylvania'),
    ('RI','Rhode Island'),
    ('SC','South Carolina'),
    ('SD','South Dakota'),
    ('TN','Tennessee'),
    ('TX','Texas'),
    ('UT','Utah'),
    ('VT','Vermont'),
    ('VA','Virginia'),
    ('WA','Washington'),
    ('WV','West Virginia'),
    ('WI','Wisconsin'),
    ('WY','Wyoming')
)

# Models

sheet_icon_path = PathAndRename("sheets/icons/")
temp_document_doc_path = PathAndRename("documents/docs/templates/")
vars_document_doc_path = PathAndRename("documents/docs/variable/")

class Document(models.Model):
    """A Document.
    """
    class Meta:
        verbose_name = 'Document'
        verbose_name_plural = 'Documents'

    objects       = DocumentManager()

    # Fields
    created_at    = models.DateTimeField(editable=False)
    updated_at    = models.DateTimeField()

    title         = models.CharField(max_length=250)

    page_count    = models.IntegerField(null=True, blank=True)
    template_doc  = models.FileField(upload_to=temp_document_doc_path,
                                     storage=PrivateMediaStorage(),
                                     null=True,
                                     blank=True)
    variable_doc  = models.FileField(upload_to=vars_document_doc_path,
                                     storage=PrivateMediaStorage(),
                                     null=True,
                                     blank=True)

    variable_dict = JSONField(null=True, blank=True)

    # The Sheet connected to the doc.
    sheet         = models.OneToOneField('Sheet',\
                                         related_name="doc",\
                                         on_delete=models.SET_NULL,
                                         null=True, blank=True)

    # Methods
    def __str__(self):
        return f'<Document — title: {self.title}>'

    def __unicode__(self):
        return f'<Document — title: {self.title}>'

    def save(self, *args, **kwargs):
        if not self.pk:
            self.created_at = timezone.now()
        self.updated_at = timezone.now()
        super(Document, self).save(*args, **kwargs)

def document_post_save(sender, instance, created, *args, **kwargs):
    """Trigger after a new instance of Document is created.
    """
    if instance.template_doc and not instance.variable_dict:
        # Open the template doc.
        r = requests.get(instance.template_doc.url)
        with open('/tmp/template.docx', 'wb') as f:
            f.write(r.content)
        doc = docx.Document('/tmp/template.docx')
        # Create the variable doc.
        variable_dict = variable_dict_from_doc(doc)
        variable_doc = new_doc_from_template_with_var_substitution(variable_dict, doc) # THIS LINE?
        variable_doc.save('/tmp/template.docx')
        reopen = open('/tmp/template.docx', 'rb')
        new_file = File(reopen)
        instance.variable_dict = variable_dict
        instance.variable_doc.save(os.path.basename(instance.template_doc.name), new_file, save=True)

# Connect the post-save receivers.
post_save.connect(document_post_save, sender=Document)

class Sheet(models.Model):
    """A sheet.
    """
    class Meta:
        verbose_name = 'sheet'
        verbose_name_plural = 'sheets'

    objects           = SheetManager()

    # Fields
    created_at        = models.DateTimeField(editable=False)
    updated_at        = models.DateTimeField(editable=False)

    title             = models.CharField(max_length=250)
    subtitle          = models.CharField(max_length=250)

    editable          = models.BooleanField(default=True)
    published         = models.BooleanField(default=False)

    icon              = models.ImageField(upload_to=sheet_icon_path,
                            null=True, blank=True)

    download_title    = models.CharField(max_length=250, null=True, blank=True)
    short_description = models.CharField(max_length=1000, null=True, blank=True)
    long_description  = models.TextField(null=True, blank=True)
    use_case          = models.CharField(max_length=50, choices=USE_CASE_CHOICES,
                                         null=True, blank=True)
    page_count        = models.IntegerField(null=True, blank=True)
    excerpt           = models.TextField(null=True, blank=True)
    bullets           = ArrayField(models.CharField(max_length=250), 
                                   default=list, null=True, blank=True)

    toc               = ArrayField(ArrayField(models.CharField(max_length=200)), default=list)

    hours_saved       = models.IntegerField(null=True, blank=True)
    traditional_cost  = models.FloatField(null=True, blank=True)

    slug              = models.SlugField(blank=True)

    # The Item connected to the sheet.
    item              = models.OneToOneField(Item,\
                                             related_name="sheet",\
                                             on_delete=models.SET_NULL,\
                                             null=True, blank=True)

    # Methods
    def __str__(self):
        return f'<Sheet — title: {self.title}>'

    def __unicode__(self):
        return f'<Sheet — title: {self.title}>'

    def save(self, *args, **kwargs):
        """Saves the Sheet.
        """
        if not self.pk:
            self.created_at = timezone.now()
        self.updated_at = timezone.now()
        self.slug = unique_slug_generator(self)
        super(Sheet, self).save(*args, **kwargs)

def create_item(sender, instance, created, *args, **kwargs):
    """Triggers an audit after a new instance of Sheet is created.
    """
    if created:
        item = Item.objects.new(sheet=instance)
    
# Connect the post-save receivers.
post_save.connect(create_item, sender=Sheet)

def user_variable_doc_path(instance, filename):
    ext = filename.split('.')[-1]
    filename = shortuuid.uuid()
    download_title = instance.doc.sheet.download_title
    return os.path.join("documents/docs/user/variable/", filename, '{}.{}'.format(download_title, ext))

def user_variable_pdf_path(instance, filename):
    ext = filename.split('.')[-1]
    filename = shortuuid.uuid()
    download_title = instance.doc.sheet.download_title
    return os.path.join("documents/pdfs/user/variable/", filename, '{}.{}'.format(download_title, ext))

class UserSheet(models.Model):
    """A company.
    """
    class Meta:
        verbose_name = 'User Sheet'
        verbose_name_plural = 'User Sheets'

    objects = UserSheetManager()

    # Fields
    created_at                        = models.DateTimeField(editable=False)
    updated_at                        = models.DateTimeField()

    total_input_count                 = models.IntegerField(default=0, null=True, blank=True)
    required_input_count              = models.IntegerField(default=0, null=True, blank=True)
    var_count                         = models.IntegerField(default=0, null=True, blank=True)
    completed_input_count             = models.IntegerField(default=0, null=True, blank=True)
    completed_required_input_count    = models.IntegerField(default=0, null=True, blank=True)

    user_variable_dict                = JSONField(null=True, blank=True)
    user_variable_doc                 = models.FileField(upload_to=user_variable_doc_path,
                                                storage=PrivateMediaStorage(),
                                                null=True,
                                                blank=True)
    
    user_variable_pdf                 = models.FileField(upload_to=user_variable_pdf_path,
                                                storage=PrivateMediaStorage(),
                                                null=True,
                                                blank=True)

    # Relationships
    user                     = models.ForeignKey(User,
                                         related_name='user_sheets',\
                                         on_delete=models.CASCADE)
    sheet                    = models.ForeignKey(Sheet,
                                         related_name='user_sheets',\
                                         on_delete=models.CASCADE)
    doc                      = models.ForeignKey(Document,
                                         related_name='user_sheets',
                                         on_delete=models.CASCADE, 
                                         null=True)
                                   
    # Methods
    def __str__(self):
        return f'<UserSheet — user: {self.user} — sheet: {self.sheet}>'

    def __unicode__(self):
        return f'<UserSheet — user: {self.user} — sheet: {self.sheet}>'

    def get_doc_download_url(self):
        return self.user_variable_doc.url

    def get_pdf_download_url(self):
        updated_string = self.user_variable_pdf.url.replace("private/", "", 1)
        return updated_string

    def save_document(self):
        doc = docx.Document(self.doc.variable_doc.open())
        variable_doc = save_doc_with_user_var_substitution(self.user_variable_dict, doc)
        variable_doc.save(f'/tmp/{self.id}.docx')
        reopen = open(f'/tmp/{self.id}.docx', 'rb')
        new_file = File(reopen)
        self.user_variable_doc.save(os.path.basename(self.doc.template_doc.name), new_file, save=True)
        self.save()

    def generate_pdf(self):
        lambda_url = 'https://fwyzic75s8.execute-api.us-east-1.amazonaws.com/default/lambda_to_pdf_converter'
        match = re.findall("variable/(.*?)\?", self.user_variable_doc.url)
        splits = match[0].split('/')
        directory = splits[0]
        filename = splits[1]
        tries = 1
        success = False
        while not success:
            try:
                response = requests.post(lambda_url, json={"directory": directory, "filename": filename})
                parsed_response = json.loads(response.content)
                parsed_response = parsed_response[8:]
                self.user_variable_pdf = parsed_response
                self.save()
                success = True
            except TypeError:
                if tries < 5:
                    tries += 1
                    continue
                else:
                    break
            
    def save(self, *args, **kwargs):
        """Saves the User Sheet.
        """
        if not self.pk:
            self.created_at = timezone.now()
        self.updated_at = timezone.now()
        super(UserSheet, self).save(*args, **kwargs)

icon_path = PathAndRename("companies/icons/")

def post_user_sheet_receiver(sender, instance, created, *args, **kwargs):
    if created is True:
        if instance.doc:
            # Open the Variable Doc.
            variable_doc = instance.doc.variable_doc
            instance.user_variable_doc.save(\
                os.path.basename(variable_doc.name), 
                variable_doc, 
                save=True)

post_save.connect(post_user_sheet_receiver, sender=UserSheet)

def post_charge_save_receiver(sender, instance, created, *args, **kwargs):
    # Create the UserSheets for each sheet in the order associated with the charge.
    if created is True:
        billing_profile = instance.billing_profile
        user = billing_profile.user
        order_items = instance.order.order_items.all()
        items = [order_item.item for order_item in order_items]
        sheets = [item.sheet for item in items]
        for sheet in sheets:
            UserSheet.objects.new(user=user, sheet=sheet)

post_save.connect(post_charge_save_receiver, sender=Charge)

class Company(models.Model):
    """A company.
    """
    class Meta:
        verbose_name = 'Company'
        verbose_name_plural = 'Companies'

    objects     = CompanyManager()

    # Fields
    created_at  = models.DateTimeField(editable=False)
    updated_at  = models.DateTimeField()

    name        = models.CharField(max_length=255)
    type_of_inc = models.CharField(max_length=4, choices=INC_CHOICES, null=True, blank=True)
    category    = models.CharField(max_length=4, choices=CATEGORY_CHOICES, null=True, blank=True)
    new_status  = models.CharField(max_length=1, choices=NEW_CHOICES, null=True, blank=True)
    use_case    = models.CharField(max_length=50, choices=USE_CASE_CHOICES, null=True, blank=True)
    state       = models.CharField(max_length=2, choices=STATE_CHOICES, null=True, blank=True)

    verified    = models.BooleanField(default=False)

    icon        = ProcessedImageField(upload_to=icon_path,
                                      null=True, blank=True,
                                      processors=[ResizeToFill(150, 150)],
                                      format='JPEG',
                                      options={'quality': 100})

    # Relationships
    user        = models.ForeignKey(User,
                                    related_name='companies',
                                    on_delete=models.CASCADE, 
                                    null=True)

    # Methods
    def __str__(self):
        return f'<Company — name: {self.name}>'

    def __unicode__(self):
        return f'<Company — name: {self.name}>'

    def save(self, *args, **kwargs):
        """Saves the Company.
        """
        if not self.pk:
            self.created_at = timezone.now()
        self.updated_at = timezone.now()
        super(Company, self).save(*args, **kwargs)

def post_company_save_receiver(sender, instance, created, *args, **kwargs):
    if instance.type_of_inc and instance.state and instance.name and instance.category and instance.use_case:
        if instance.verified is False:
            instance.verified = True
            instance.save()
    else: 
        if instance.verified is True:
            instance.verified = False
            instance.save()

post_save.connect(post_company_save_receiver, sender=Company)